import os
import re
import time
import nltk
import pickle
import logging
import unicodedata
import pymorphy2

from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table, _Row
from docx.text.paragraph import Paragraph
from nltk.corpus.reader.api import CorpusReader, CategorizedCorpusReader
from nltk import pos_tag, sent_tokenize, wordpunct_tokenize


CAT_PATTERN = r'([a-z_\s]+)/.*'
DOC_PATTERN = r'(?!\.)\w+\.docx'
PKL_PATTERN = r'(?!\.)[\w_\s]+/[\w_\s]+\.pickle' 
DIAGNOS_PATTERN = 'Клинический диагноз'   
THERAPY_PATTERN = 'Проведено лечение'

CAT_DICT = {1.0: 'cat_bronkhit',
            5.0: 'cat_astma',
            6.0: 'cat_mukovistsidoz',
            7.0: 'cat_plevryt',
            8.0: 'cat_pnevmoniya',
            0.0: 'cat_nan',
            }


class DOCXCorpusReader(CorpusReader):
    """
    A corpus reader for raw HTML documents to enable preprocessing.
    """

    def __init__(self, root, fileids=DOC_PATTERN, encoding='utf8'):
        """
        Initialize the corpus reader.  Categorization arguments
        (``cat_pattern``, ``cat_map``, and ``cat_file``) are passed to
        the ``CategorizedCorpusReader`` constructor.  The remaining
        arguments are passed to the ``CorpusReader`` constructor.
        """

        # Initialize the NLTK corpus reader objects
        CorpusReader.__init__(self, root, fileids, encoding)


    def docs(self, fileids=None, categories=None):
        """
        Returns the complete text of an HTML document, closing the document
        after we are done reading it and yielding it in a memory safe fashion.
        """

        # Create a generator, loading one document into memory at a time.
        for path in self.abspaths(fileids):
            yield Document(path)

    def sizes(self, fileids=None, categories=None):
        """
        Returns a list of tuples, the fileid and size on disk of the file.
        This function is used to detect oddly large files in the corpus.
        """
        # Resolve the fileids and the categories

        # Create a generator, getting every path and computing filesize
        for path in self.abspaths(self.fileids()):
            yield os.path.getsize(path)

    def iter_block_items_(self, parent):
        """
        Generate a reference to each paragraph and table child within *parent*,
        in document order. Each returned value is an instance of either Table or
        Paragraph. *parent* would most commonly be a reference to a main
        Document object, but also works for a _Cell object, which itself can
        contain paragraphs and tables.
        """
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        elif isinstance(parent, _Row):
            parent_elm = parent._tr
        else:
            raise ValueError("something's not right")
            
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def paras(self, fileids=None, categories=None):
        """
        Uses BeautifulSoup to parse the paragraphs from the HTML.
        """
        for doc in self.docs(fileids):
            table_id = 0
            for block in self.iter_block_items_(doc):
                paragr = ''
                #print(block.text if isinstance(block, Paragraph) else '<table>')
                if isinstance(block, Paragraph):
                    if len(re.sub(r'\s+', '', block.text)) == 0: 
                        continue
                    else:
                        paragr = block.text
                elif isinstance(block, Table):
                    paragr = f'table_{table_id}'
                    table_id += 1
                yield paragr

        # print("\t".join(table_header))
        #     for idx, para in enumerate(doc.paragraphs):
        #         if idx < 3: continue
        #     #     text = re.sub(r'[\d]+[\/\.]\d+[\/\.]*\d*', '', para.text)
        #     #     # text = re.sub(r'Ф\.И\.О\.', 'ФИО', text)
        #     #     if DIAGNOS_PATTERN in text: continue
        #         if len(re.sub(r'\s+', '', para.text)) == 0: continue
        #     #     if THERAPY_PATTERN in text: break
        #     #     yield text
        #         yield para.text

    def tables(self, fileids=None, categoties=None):
        for doc in self.docs(fileids):
            for table in doc.tables:
                yield [ 
                    [
                        re.sub(r'\s+', '', cell.text)
                        for cell in row.cells
                    ]
                    for row in table.rows
                ]

    # def table_headers(self, fileids=None, categories=None):
    #     for table in self.tables(fileids, categories):
    #         for cell in table.rows[0].cells:
    #             yield cell.text
     
    def sents(self, fileids=None, categories=None):
        """
        Uses the built in sentence tokenizer to extract sentences from the
        paragraphs. Note that this method uses BeautifulSoup to parse HTML.
        """
        for paragraph in self.paras(fileids, categories):
            for sentence in sent_tokenize(paragraph, language='russian'):
                yield sentence

    def words(self, fileids=None, categories=None):
        """
        Uses the built in word tokenizer to extract tokens from sentences.
        Note that this method uses BeautifulSoup to parse HTML content.
        """
        for sentence in self.sents(fileids, categories):
            for token in wordpunct_tokenize(sentence):
                yield token

    def tokenize(self, fileids=None, categories=None):
        """
        Segments, tokenizes, and tags a document in the corpus.
        """
        morph = pymorphy2.MorphAnalyzer()
        for paragraph in self.paras(fileids=fileids):
            yield [
                (word, morph.parse(word)[0].tag.POS)
                # pos_tag(wordpunct_tokenize(sent), lang='rus')
                for sent in sent_tokenize(paragraph)
                for word in wordpunct_tokenize(sent)
            ]

    def describe(self, fileids=None, categories=None):
        """
        Performs a single pass of the corpus and
        returns a dictionary with a variety of metrics
        concerning the state of the corpus.
        """
        started = time.time()

        # Structures to perform counting.
        counts  = nltk.FreqDist()
        tokens  = nltk.FreqDist()

        # Perform single pass over paragraphs, tokenize and count
        for para in self.paras(fileids, categories):
            counts['paras'] += 1

            for sent in sent_tokenize(para):
                counts['sents'] += 1

                for word in wordpunct_tokenize(sent):
                    counts['words'] += 1
                    tokens[word] += 1

        # Compute the number of files and categories in the corpus
        n_fileids = len(self.fileids())
        # n_topics  = len(self.categories(self.resolve(fileids, categories)))

        # Return data structure with information
        return {
            'files':  n_fileids,
            # 'topics': n_topics,
            'paras':  counts['paras'],
            'sents':  counts['sents'],
            'words':  counts['words'],
            'vocab':  len(tokens),
            'lexdiv': float(counts['words']) / float(len(tokens)),
            'ppdoc':  float(counts['paras']) / float(n_fileids),
            'sppar':  float(counts['sents']) / float(counts['paras']),
            'secs':   time.time() - started,
        }

class Preprocessor(object):
    """
    The preprocessor wraps a corpus object (usually a `HTMLCorpusReader`)
    and manages the stateful tokenization and part of speech tagging into a
    directory that is stored in a format that can be read by the
    `HTMLPickledCorpusReader`. This format is more compact and necessarily
    removes a variety of fields from the document that are stored in the JSON
    representation dumped from the Mongo database. This format however is more
    easily accessed for common parsing activity.
    """

    def __init__(self, corpus, target=None, **kwargs):
        """
        The corpus is the `HTMLCorpusReader` to preprocess and pickle.
        The target is the directory on disk to output the pickled corpus to.
        """
        self.corpus = corpus
        # with open('/home/igor/Development/atap/data/dummy.pkl', 'rb') as f:
        #     target = pickle.load(f).to_dict()
        self.target = target
        self.morph = pymorphy2.MorphAnalyzer()
        self.stopwords  = nltk.corpus.stopwords.words('russian')
        self.stopwords.extend(['анализ', 'год', 'плата', 'адрес', 'поступить', 'выписать' 'что', 'это', 'так', 'вот', 'быть', 'как', 'в', '—', '–', 'к', 'на', '...'])

    def fileids(self, fileids=None, categories=None):
        """
        Helper function access the fileids of the corpus
        """
        # fileids = self.corpus.resolve(fileids, categories)
        # if fileids:
        #     return fileids
        return self.corpus.fileids()

    def abspath(self, fileid):
        """
        Returns the absolute path to the target fileid from the corpus fileid.
        """
        # Find the directory, relative from the corpus root.
        name = fileid.split('.')[0]
        category = fileid.split('_')[0]
        # Create the pickle file extension
        basename  = name + '.pickle'

        # Return the path to the file relative to the target.
        return os.path.normpath(os.path.join(self.target, category, basename))

    def is_punct(self, token):
        return all(
            unicodedata.category(char).startswith('P') for char in token
        )

    def is_stopword(self, token):
        return token.lower() in self.stopwords

    def lemmatize(self, token):
        return self.morph.parse(token)[0].normal_form

    def tokenize(self, fileid):
        """
        Segments, tokenizes, and tags a document in the corpus. Returns a
        generator of paragraphs, which are lists of sentences, which in turn
        are lists of part of speech tagged words.
        """
        for paragraph in self.corpus.paras(fileids=fileid):
            sents = []
            for sent in sent_tokenize(paragraph, language='russian'):
                words = []
                for word in wordpunct_tokenize(sent):
                    token = self.lemmatize(word)
                    if not self.is_punct(token) and not self.is_stopword(token):

                        words.append((token, str(self.morph.parse(word)[0].tag.POS)))

                sents.append(words)
            yield sents
            # yield [
            #     (word, morph.parse(word)[0].tag.POS)
            #     # pos_tag(wordpunct_tokenize(sent), lang='rus')
            #     for sent in sent_tokenize(paragraph, language='russian')
            #     for word in wordpunct_tokenize(sent)
            # ]
            # yield [
            #     pos_tag(wordpunct_tokenize(sent), lang='rus')
            #     for sent in sent_tokenize(paragraph, language='russian')
            # ]

    def process(self, fileid):
        """
        For a single file does the following preprocessing work:
            1. Checks the location on disk to make sure no errors occur.
            2. Gets all paragraphs for the given text.
            3. Segements the paragraphs with the sent_tokenizer
            4. Tokenizes the sentences with the wordpunct_tokenizer
            5. Tags the sentences using the default pos_tagger
            6. Writes the document as a pickle to the target location.
        This method is called multiple times from the transform runner.
        """
        # Compute the outpath to write the file to.
        target = self.abspath(fileid)
        parent = os.path.dirname(target)

        # Make sure the directory exists
        if not os.path.exists(parent):
            os.makedirs(parent)

        # Make sure that the parent is a directory and not a file
        if not os.path.isdir(parent):
            raise ValueError(
                "Please supply a directory to write preprocessed data to."
            )
        document = {}
        # Create a data structure for the pickle
        document['text'] = list(self.corpus.paras(fileid))
        document['tables'] = list(self.corpus.tables(fileid))
        # print(document)
        # Open and serialize the pickle to disk
        with open(target, 'wb') as f:
            pickle.dump(document, f, pickle.HIGHEST_PROTOCOL)

        # Clean up the document
        del document

        # Return the target fileid
        return target

    def transform(self, fileids=None, categories=None):
        """
        Transform the wrapped corpus, writing out the segmented, tokenized,
        and part of speech tagged corpus as a pickle to the target directory.
        This method will also directly copy files that are in the corpus.root
        directory that are not matched by the corpus.fileids().
        """
        # Make the target directory if it doesn't already exist
        if not os.path.exists(self.target):
            os.makedirs(self.target)

        # Resolve the fileids to start processing and return the list of 
        # target file ids to pass to downstream transformers. 
        return [
            self.process(fileid)
            for fileid in self.fileids(fileids, categories)
        ]